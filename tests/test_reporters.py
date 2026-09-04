# tests/test_reporters.py
import unittest
import shutil
from pathlib import Path
import docx
import openpyxl

from app.services.chart_generator import ChartGenerator
from app.services.excel_reporter import ExcelReporter
from app.services.word_reporter import WordReporter


class TestReporters(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.test_dir = Path("scratch/test_reports_unit")
        cls.test_dir.mkdir(parents=True, exist_ok=True)

        cls.sample_analisis_ia = {
            "nivel_alerta": "MEDIO",
            "score_relevancia": 7.8,
            "tipo_portal": "Portal Universitario / Convocatorias",
            "resumen_ejecutivo": "Auditoría de prueba realizada sobre publicaciones académicas.",
            "analisis_evolucion": "Línea base registrada sin versiones anteriores.",
            "puntos_atencion_urgentes": ["Fecha límite de matrícula en 3 días."],
            "categorias": [
                {"nombre": "Académico", "cantidad": 5, "porcentaje": 60},
                {"nombre": "Investigación", "cantidad": 3, "porcentaje": 40}
            ],
            "sentimientos": {"positivo": 4, "neutro": 3, "negativo": 1},
            "entidades_estructuradas": {
                "instituciones_y_empresas": ["Universidad Nacional", "Consejo de Facultad"],
                "personas_relevantes": ["Dr. Decano"],
                "marcos_legales_o_normas": ["Estatuto Universitario"]
            },
            "entidades_clave": ["Universidad Nacional"],
            "conclusiones": ["Cronograma académico en curso."],
            "recomendaciones_estrategicas": ["Verificar publicación de horarios."]
        }

        cls.sample_articulos = [
            {"titulo": "Convocatoria a concurso docente 2026", "url": "https://uni.edu.pe/concurso-1"},
            {"titulo": "Calendario de matrícula extemporánea", "url": "https://uni.edu.pe/matricula-2"}
        ]

        cls.sample_delta = {
            "es_lista": True,
            "total_anteriores": 2,
            "total_actuales": 2,
            "total_nuevos": 1,
            "total_salientes": 0,
            "total_mantenidos": 1,
            "tasa_rotacion_pct": 50.0,
            "nuevos_articulos": [cls.sample_articulos[0]]
        }

    @classmethod
    def tearDownClass(cls):
        if cls.test_dir.exists():
            shutil.rmtree(cls.test_dir, ignore_errors=True)

    def test_chart_generator(self):
        """Verifica la generación de los 3 tipos de gráficos estadísticos."""
        c_cat = ChartGenerator.generar_grafico_categorias(self.sample_analisis_ia["categorias"], self.test_dir)
        c_sent = ChartGenerator.generar_grafico_sentimientos(self.sample_analisis_ia["sentimientos"], self.test_dir)
        c_delta = ChartGenerator.generar_grafico_delta(self.sample_delta, self.test_dir)

        for c in [c_cat, c_sent, c_delta]:
            self.assertTrue(c.exists())
            self.assertGreater(c.stat().st_size, 0)

    def test_word_reporter_generation(self):
        """Verifica la generación íntegra y lectura de un informe formal en Word (.docx)."""
        word_out = self.test_dir / "test_doc.docx"
        WordReporter.generar_reporte(
            url="https://uni.edu.pe",
            site_title="Portal Universitario",
            analisis_ia=self.sample_analisis_ia,
            articulos=self.sample_articulos,
            delta=self.sample_delta,
            chart_paths=[],
            output_file=word_out
        )

        self.assertTrue(word_out.exists())
        self.assertGreater(word_out.stat().st_size, 1000)

        # Validar que python-docx pueda abrirlo sin errores de XML
        doc = docx.Document(str(word_out))
        self.assertGreater(len(doc.paragraphs), 5)
        self.assertGreater(len(doc.tables), 2)

    def test_excel_reporter_generation(self):
        """Verifica la generación y presencia de las 3 hojas en el libro Excel (.xlsx)."""
        excel_out = self.test_dir / "test_book.xlsx"
        ExcelReporter.generar_reporte(
            url="https://uni.edu.pe",
            site_title="Portal Universitario",
            analisis_ia=self.sample_analisis_ia,
            articulos=self.sample_articulos,
            delta=self.sample_delta,
            chart_paths=[],
            output_file=excel_out
        )

        self.assertTrue(excel_out.exists())
        self.assertGreater(excel_out.stat().st_size, 1000)

        wb = openpyxl.load_workbook(str(excel_out))
        self.assertIn("Dashboard Ejecutivo", wb.sheetnames)
        self.assertIn("Detalle de Contenidos", wb.sheetnames)
        self.assertIn("Matriz de Inteligencia", wb.sheetnames)


if __name__ == "__main__":
    unittest.main()
