# app/services/word_reporter.py
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls


def set_cell_background(cell, fill_hex: str):
    """Aplica color de fondo a una celda de tabla en Word."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Configura padding interno para celdas."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tcMar)


def add_callout_box(doc, texto: str, titulo: Optional[str] = None, border_hex="2563EB", bg_hex="F8FAFC"):
    """Crea una caja destacada (Callout box) con borde izquierdo grueso y fondo suave."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)

    # Borde izquierdo grueso
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/>'
        f'<w:top w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(borders)

    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(4)

    if titulo:
        r_tit = p.add_run(f"{titulo}\n")
        r_tit.font.bold = True
        r_tit.font.size = Pt(10.5)
        r_tit.font.color.rgb = RGBColor(15, 23, 42)

    r_txt = p.add_run(texto)
    r_txt.font.size = Pt(9.5)
    r_txt.font.color.rgb = RGBColor(51, 65, 85)


class WordReporter:
    @staticmethod
    def generar_reporte(
        url: str,
        site_title: str,
        analisis_ia: Dict[str, Any],
        articulos: List[Dict[str, Any]],
        delta: Optional[Dict[str, Any]],
        chart_paths: List[Path],
        output_file: Path
    ) -> Path:
        """
        Genera un informe formal en Microsoft Word (.docx) con membrete ejecutivo,
        tarjetas KPI, cajas destacadas, gráficos y tablas con diseño editorial.
        """
        doc = docx.Document()

        # Configurar márgenes de 0.8 pulgadas
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Paleta institucional
        c_navy = RGBColor(15, 23, 42)      # Slate 900
        c_blue = RGBColor(37, 99, 235)     # Royal Blue
        c_gray = RGBColor(100, 116, 139)   # Slate 500

        # Nivel de alerta y color correspondiente
        nivel_alerta = str(analisis_ia.get("nivel_alerta", "MEDIO")).upper()
        color_alerta_hex = "10B981" if nivel_alerta == "BAJO" else ("F59E0B" if nivel_alerta == "MEDIO" else "EF4444")

        # ---------------------------------------------------------
        # 1. ENCABEZADO Y MEMBRETE EJECUTIVO
        # ---------------------------------------------------------
        p_pre = doc.add_paragraph()
        r_pre = p_pre.add_run("SISTEMA DE INTELIGENCIA Y AUDITORÍA DE CONTENIDOS WEB")
        r_pre.font.size = Pt(9)
        r_pre.font.bold = True
        r_pre.font.color.rgb = c_blue
        p_pre.paragraph_format.space_after = Pt(2)

        p_main = doc.add_paragraph()
        r_main = p_main.add_run("INFORME EJECUTIVO DE MONITOREO")
        r_main.font.size = Pt(22)
        r_main.font.bold = True
        r_main.font.color.rgb = c_navy
        p_main.paragraph_format.space_after = Pt(4)

        p_sub = doc.add_paragraph()
        r_sub = p_sub.add_run(f"Auditoría de Contenidos: {site_title}")
        r_sub.font.size = Pt(13)
        r_sub.font.color.rgb = c_gray
        p_sub.paragraph_format.space_after = Pt(14)

        # ---------------------------------------------------------
        # 2. TARJETAS MÉTRICAS DE IMPACTO (KPIS)
        # ---------------------------------------------------------
        total_items = len(articulos) if isinstance(articulos, list) else 1
        total_nuevos = delta.get("total_nuevos", 0) if delta else 0
        rotacion = delta.get("tasa_rotacion_pct", 0) if delta else 0.0

        t_kpis = doc.add_table(rows=1, cols=4)
        t_kpis.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_kpis.autofit = False

        datos_kpis = [
            ("TOTAL CONTENIDOS", str(total_items), "Elementos indexados", "F1F5F9", "1E293B"),
            ("NOVEDADES", f"+{total_nuevos}", "Nuevos registros", "ECFDF5", "059669"),
            ("ROTACIÓN", f"{rotacion}%", "Tasa de recambio", "F0F9FF", "0284C7"),
            ("NIVEL DE ALERTA", nivel_alerta, f"Score: {analisis_ia.get('score_relevancia', 7.5)}/10", "FEF2F2" if nivel_alerta in ("ALTO", "CRÍTICO") else "FFFBEB", color_alerta_hex)
        ]

        for i, (titulo_kpi, valor_kpi, subtitulo_kpi, bg_col, text_col) in enumerate(datos_kpis):
            cell = t_kpis.cell(0, i)
            cell.width = Inches(1.7)
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)

            # Borde sutil
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_pr.append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:bottom w:val="single" w:sz="18" w:color="{text_col}"/></w:tcBorders>'))

            p_cell = cell.paragraphs[0]
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cell.paragraph_format.space_after = Pt(2)

            r_kpi_tit = p_cell.add_run(f"{titulo_kpi}\n")
            r_kpi_tit.font.size = Pt(7.5)
            r_kpi_tit.font.bold = True
            r_kpi_tit.font.color.rgb = RGBColor(100, 116, 139)

            r_kpi_val = p_cell.add_run(f"{valor_kpi}\n")
            r_kpi_val.font.size = Pt(15)
            r_kpi_val.font.bold = True
            r_kpi_val.font.color.rgb = RGBColor(15, 23, 42)

            r_kpi_sub = p_cell.add_run(subtitulo_kpi)
            r_kpi_sub.font.size = Pt(7.5)
            r_kpi_sub.font.color.rgb = RGBColor(148, 163, 184)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # ---------------------------------------------------------
        # 3. FICHA TÉCNICA DE AUDITORÍA
        # ---------------------------------------------------------
        t_meta = doc.add_table(rows=3, cols=2)
        t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_meta.autofit = False

        metas = [
            ("URL Objetivo:", url),
            ("Fecha y Hora de Auditoría:", datetime.now().strftime("%d/%m/%Y a las %H:%M:%S")),
            ("Clasificación de Portal:", analisis_ia.get("tipo_portal", "Portal Web Público / Institucional"))
        ]

        for idx, (label, val) in enumerate(metas):
            cell_lbl = t_meta.cell(idx, 0)
            cell_val = t_meta.cell(idx, 1)
            cell_lbl.width = Inches(2.2)
            cell_val.width = Inches(4.6)

            p1 = cell_lbl.paragraphs[0]
            r1 = p1.add_run(label)
            r1.font.bold = True
            r1.font.size = Pt(9)
            r1.font.color.rgb = c_navy

            p2 = cell_val.paragraphs[0]
            r2 = p2.add_run(val)
            r2.font.size = Pt(9)
            r2.font.color.rgb = c_gray

            set_cell_background(cell_lbl, "F8FAFC")
            set_cell_background(cell_val, "FFFFFF")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

        # ---------------------------------------------------------
        # 4. RESUMEN EJECUTIVO Y EVOLUCIÓN
        # ---------------------------------------------------------
        h1 = doc.add_heading(level=1)
        r_h1 = h1.add_run("1. Resumen Ejecutivo de Inteligencia")
        r_h1.font.color.rgb = c_navy
        r_h1.font.size = Pt(13)
        r_h1.font.bold = True

        resumen_texto = analisis_ia.get("resumen_ejecutivo", "")
        add_callout_box(doc, resumen_texto, titulo="SÍNTESIS ESTRATÉGICA", border_hex="2563EB", bg_hex="F8FAFC")

        # Puntos de atención urgentes si existen
        puntos_urgentes = analisis_ia.get("puntos_atencion_urgentes", [])
        if puntos_urgentes:
            texto_puntos = "\n".join([f"• {p}" for p in puntos_urgentes])
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            add_callout_box(doc, texto_puntos, titulo="⚠️ PUNTOS DE ATENCIÓN URGENTES Y FECHAS CRÍTICAS", border_hex="F59E0B", bg_hex="FFFBEB")

        # Evolución vs versión anterior
        p_evol = doc.add_paragraph()
        p_evol.paragraph_format.space_before = Pt(8)
        r_evol_tit = p_evol.add_run("Evolución respecto a la versión anterior:\n")
        r_evol_tit.font.bold = True
        r_evol_tit.font.size = Pt(10)
        r_evol_tit.font.color.rgb = c_navy

        r_evol_txt = p_evol.add_run(analisis_ia.get("analisis_evolucion", ""))
        r_evol_txt.font.size = Pt(9.5)
        r_evol_txt.font.color.rgb = c_gray

        # ---------------------------------------------------------
        # 5. VISUALIZACIÓN DE MÉTRICAS Y GRÁFICOS
        # ---------------------------------------------------------
        if chart_paths:
            h2 = doc.add_heading(level=1)
            r_h2 = h2.add_run("2. Análisis Gráfico y Distribución Temática")
            r_h2.font.color.rgb = c_navy
            r_h2.font.size = Pt(13)
            r_h2.font.bold = True

            for cp in chart_paths:
                if cp.exists():
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_after = Pt(10)
                    doc.add_picture(str(cp), width=Inches(5.6))

        # ---------------------------------------------------------
        # 6. MATRIZ DE ENTIDADES Y ACTORES CLAVE
        # ---------------------------------------------------------
        entidades_est = analisis_ia.get("entidades_estructuradas", {})
        h3 = doc.add_heading(level=1)
        r_h3 = h3.add_run("3. Matriz de Actores, Entidades y Normativas")
        r_h3.font.color.rgb = c_navy
        r_h3.font.size = Pt(13)
        r_h3.font.bold = True

        t_ent = doc.add_table(rows=1, cols=3)
        t_ent.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_ent.autofit = False

        headers_ent = ["Instituciones y Organismos", "Personas y Autoridades", "Normativas y Decretos"]
        for j, h in enumerate(headers_ent):
            c_hdr = t_ent.cell(0, j)
            c_hdr.width = Inches(2.26)
            set_cell_background(c_hdr, "1E293B")
            set_cell_margins(c_hdr, top=80, bottom=80, left=80, right=80)
            p_hdr = c_hdr.paragraphs[0]
            r_hdr = p_hdr.add_run(h)
            r_hdr.font.bold = True
            r_hdr.font.size = Pt(8.5)
            r_hdr.font.color.rgb = RGBColor(255, 255, 255)

        row_ent = t_ent.add_row()
        claves_ent = ["instituciones_y_empresas", "personas_relevantes", "marcos_legales_o_normas"]
        for j, k in enumerate(claves_ent):
            c_data = row_ent.cells[j]
            c_data.width = Inches(2.26)
            set_cell_background(c_data, "F8FAFC")
            set_cell_margins(c_data, top=80, bottom=80, left=80, right=80)
            items_k = entidades_est.get(k, [])
            p_data = c_data.paragraphs[0]
            p_data.paragraph_format.line_spacing = 1.2
            if items_k:
                p_data.add_run("\n".join([f"• {x}" for x in items_k[:8]])).font.size = Pt(8.5)
            else:
                r_vacio = p_data.add_run("Sin menciones explícitas.")
                r_vacio.font.size = Pt(8.5)
                r_vacio.font.italic = True
                r_vacio.font.color.rgb = RGBColor(148, 163, 184)

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

        # ---------------------------------------------------------
        # 7. TABLA DETALLADA DE CONTENIDOS Y NOVEDADES
        # ---------------------------------------------------------
        h4 = doc.add_heading(level=1)
        r_h4 = h4.add_run("4. Detalle de Contenidos Indexados y Novedades")
        r_h4.font.color.rgb = c_navy
        r_h4.font.size = Pt(13)
        r_h4.font.bold = True

        t_art = doc.add_table(rows=1, cols=4)
        t_art.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_art.autofit = False

        col_configs = [("N°", Inches(0.5)), ("Título o Asunto", Inches(3.8)), ("Estado", Inches(1.0)), ("Enlace", Inches(1.5))]
        for col_idx, (col_name, col_width) in enumerate(col_configs):
            cell_h = t_art.cell(0, col_idx)
            cell_h.width = col_width
            set_cell_background(cell_h, "1E3A8A")
            set_cell_margins(cell_h, top=80, bottom=80, left=60, right=60)
            p_h = cell_h.paragraphs[0]
            r_h = p_h.add_run(col_name)
            r_h.font.bold = True
            r_h.font.size = Pt(8.5)
            r_h.font.color.rgb = RGBColor(255, 255, 255)

        nuevas_urls = set()
        if delta and delta.get("nuevos_articulos"):
            nuevas_urls = {item.get("url") for item in delta.get("nuevos_articulos") if isinstance(item, dict)}

        lista_art = articulos[:40] if isinstance(articulos, list) else [{"titulo": site_title, "url": url}]
        for idx, art in enumerate(lista_art, start=1):
            row = t_art.add_row()
            art_url = art.get("url", "")
            es_novedad = art_url in nuevas_urls

            bg_row = "ECFDF5" if es_novedad else ("F8FAFC" if idx % 2 == 0 else "FFFFFF")

            # Col 0: N°
            c0 = row.cells[0]
            c0.width = Inches(0.5)
            set_cell_background(c0, bg_row)
            p0 = c0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r0 = p0.add_run(str(idx))
            r0.font.size = Pt(8.5)

            # Col 1: Título
            c1 = row.cells[1]
            c1.width = Inches(3.8)
            set_cell_background(c1, bg_row)
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(art.get("titulo", "Sin título")[:120])
            r1.font.size = Pt(8.5)
            r1.font.bold = es_novedad

            # Col 2: Estado
            c2 = row.cells[2]
            c2.width = Inches(1.0)
            set_cell_background(c2, bg_row)
            p2 = c2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run("Novedad" if es_novedad else "Activo")
            r2.font.size = Pt(8)
            r2.font.bold = True
            r2.font.color.rgb = RGBColor(5, 150, 105) if es_novedad else RGBColor(100, 116, 139)

            # Col 3: Enlace
            c3 = row.cells[3]
            c3.width = Inches(1.5)
            set_cell_background(c3, bg_row)
            p3 = c3.paragraphs[0]
            if art_url and art_url.startswith("http"):
                r3 = p3.add_run("Ver publicación")
                r3.font.size = Pt(8)
                r3.font.color.rgb = c_blue
                r3.font.underline = True
            else:
                p3.add_run("-").font.size = Pt(8)

            for cell_item in [c0, c1, c2, c3]:
                set_cell_margins(cell_item, top=50, bottom=50, left=60, right=60)

        # ---------------------------------------------------------
        # 8. RECOMENDACIONES ESTRATÉGICAS Y ACCIONES
        # ---------------------------------------------------------
        recs = analisis_ia.get("recomendaciones_estrategicas", [])
        if recs:
            doc.add_paragraph().paragraph_format.space_after = Pt(8)
            h5 = doc.add_heading(level=1)
            r_h5 = h5.add_run("5. Recomendaciones Estratégicas y Próximos Pasos")
            r_h5.font.color.rgb = c_navy
            r_h5.font.size = Pt(13)
            r_h5.font.bold = True

            texto_recs = "\n".join([f"✔  {r}" for r in recs])
            add_callout_box(doc, texto_recs, titulo="ACCIONES SUGERIDAS", border_hex="10B981", bg_hex="F0FDF4")

        # Guardar documento
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file))
        return output_file
